"""
DOCX Parser Module
Extracts text from Word documents (.docx files)
"""

import io
from typing import Optional

# Optional import with helpful error message
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class DOCXParser:
    """Parser for extracting text from DOCX files"""
    
    def extract_text(self, file_content: bytes, filename: Optional[str] = None) -> str:
        """
        Extract text from DOCX file content
        
        Args:
            file_content: DOCX file content as bytes
            filename: Optional filename for error messages
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If DOCX cannot be parsed
        """
        if not DOCX_AVAILABLE:
            error_msg = "python-docx is not installed. Install it with: pip install python-docx"
            if filename:
                error_msg += f"\nFile: {filename}"
            raise ImportError(error_msg)
        
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            error_msg = f"Failed to parse DOCX file"
            if filename:
                error_msg += f": {filename}"
            error_msg += f"\nError: {str(e)}"
            raise ValueError(error_msg)
    
    def is_valid_docx(self, file_content: bytes) -> bool:
        """Check if file content is a valid DOCX"""
        if not DOCX_AVAILABLE:
            return False
        try:
            doc = Document(io.BytesIO(file_content))
            return len(doc.paragraphs) > 0 or len(doc.tables) > 0
        except:
            return False

